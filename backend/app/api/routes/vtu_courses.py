import io
import re
from fastapi import APIRouter, UploadFile, File, HTTPException, status
from pydantic import BaseModel
import fitz  # PyMuPDF

router = APIRouter(
    prefix="/vtu",
    tags=["VTU Courses & Scheme Parser"],
)

# Pre-fetched official VTU B.E. Degree Courses
PREFETCHED_VTU_COURSES = [
    {"code": "CSE", "name": "Computer Science & Engineering", "is_vtu_standard": True},
    {"code": "CSE-AIML", "name": "CSE (Artificial Intelligence & Machine Learning)", "is_vtu_standard": True},
    {"code": "CSE-DS", "name": "CSE (Data Science)", "is_vtu_standard": True},
    {"code": "ECE", "name": "Electronics & Communication Engineering", "is_vtu_standard": True},
    {"code": "EEE", "name": "Electrical & Electronics Engineering", "is_vtu_standard": True},
    {"code": "ISE", "name": "Information Science & Engineering", "is_vtu_standard": True},
    {"code": "AI&DS", "name": "Artificial Intelligence & Data Science", "is_vtu_standard": True},
    {"code": "ME", "name": "Mechanical Engineering", "is_vtu_standard": True},
    {"code": "CIV", "name": "Civil Engineering", "is_vtu_standard": True},
    {"code": "CH", "name": "Chemical Engineering", "is_vtu_standard": True},
    {"code": "BME", "name": "Biomedical Engineering", "is_vtu_standard": True},
]


class VTUSubject(BaseModel):
    code: str
    name: str
    department: str = "Computer Science & Engineering"
    category: str  # "theory" or "practical"
    weekly_hours: int = 4


def infer_department(code: str, name: str, default_course: str = "CSE") -> str:
    c = (code or "").upper().strip()
    n = (name or "").upper().strip()

    # 1. Basic Sciences & Humanities
    if any(k in c for k in ["MAT", "MTH"]) or any(k in n for k in ["MATH", "STATISTICS", "PROBABILITY", "CALCULUS", "LINEAR ALGEBRA", "NUMERICAL", "DISCRETE"]):
        return "Mathematics"
    if "PHY" in c or "PHYSICS" in n or "QUANTUM" in n:
        return "Physics"
    if "CHE" in c or "CHEMISTRY" in n:
        return "Chemistry"
    if any(k in c for k in ["HUM", "ENG", "CIP", "KAN", "IDT", "SFH"]) or any(k in n for k in ["CONSTITUTION", "ENVIRONMENT", "MANAGEMENT", "ENTREPRENEURSHIP", "ENGLISH", "KANNADA", "ETHICS", "YOGA", "COMMUNITY"]):
        return "Humanities & Social Sciences"

    # 2. Biological / Biomedical
    if any(k in c for k in ["BBM", "BM", "BT", "BIO"]) or any(k in n for k in ["BIOMEDICAL", "BIOLOGY", "BIOMECHANICS", "BIODYNAMICS", "BIOPROCESS", "GENETICS"]):
        return "Biomedical Engineering"

    # 3. Electrical & Electronics
    if any(k in c for k in ["BEE", "EE"]) or any(k in n for k in ["SWITCHGEAR", "POWER SYSTEM", "ELECTRICAL", "HIGH VOLTAGE", "DRIVES", "TRANSFORMER"]):
        return "Electrical & Electronics Engineering"

    # 4. Electronics & Communication
    if any(k in c for k in ["BEC", "EC"]) or any(k in n for k in ["ANTENNA", "MICROWAVE", "WIRELESS", "COMMUNICATION", "ANALOG", "VLSI", "EMBEDDED", "SIGNAL PROCESSING", "DSP", "MICROCONTROLLER", "DIGITAL DESIGN"]):
        return "Electronics & Communication Engineering"

    # 5. Mechanical Engineering
    if any(k in c for k in ["BME", "ME"]) or any(k in n for k in ["FINITE ELEMENT", "HYDRAULIC", "PNEUMATIC", "THERMODYNAMIC", "MECHANICS", "TURBOMACHINERY", "ROBOTICS", "HEAT TRANSFER", "MANUFACTURING", "AUTOMOBILE"]):
        return "Mechanical Engineering"

    # 6. Civil Engineering
    if any(k in c for k in ["BCV", "CV", "CIV"]) or any(k in n for k in ["STEEL STRUCTURE", "CONCRETE", "SURVEYING", "GEOTECHNICAL", "STRUCTURAL", "ENVIRONMENTAL ENG", "HYDROLOGY"]):
        return "Civil Engineering"

    # 7. Chemical Engineering
    if any(k in c for k in ["BCH", "CH"]) or any(k in n for k in ["PROCESS MODELING", "PROCESS CONTROL", "HEAT TRANSFER", "MASS TRANSFER", "REACTION ENG"]):
        return "Chemical Engineering"

    # 8. AI & Data Science
    if any(k in c for k in ["BAI", "BCD", "BDS", "BAD", "AD"]) or any(k in n for k in ["DEEP LEARNING", "ARTIFICIAL INTELLIGENCE", "MACHINE LEARNING", "NEURAL NETWORK", "NATURAL LANGUAGE", "COMPUTER VISION"]):
        return "Artificial Intelligence & Machine Learning"

    # 9. Information Science
    if any(k in c for k in ["BIS", "IS"]) or "INFORMATION SCIENCE" in n:
        return "Information Science & Engineering"

    # 10. Computer Science
    if any(k in c for k in ["BCS", "CS", "CSE"]) or any(k in n for k in ["DATA STRUCTURE", "ALGORITHM", "OPERATING SYSTEM", "DATABASE", "JAVA", "PYTHON", "C++", "COMPILER", "CLOUD", "SOFTWARE", "CYBER", "WEB", "NETWORK"]):
        return "Computer Science & Engineering"

    # Course-code fallback mapping
    course_dept_map = {
        "CSE": "Computer Science & Engineering",
        "CSE-AIML": "Artificial Intelligence & Machine Learning",
        "CSE-DS": "Artificial Intelligence & Machine Learning",
        "ISE": "Information Science & Engineering",
        "AI&DS": "Artificial Intelligence & Data Science",
        "ECE": "Electronics & Communication Engineering",
        "EEE": "Electrical & Electronics Engineering",
        "ME": "Mechanical Engineering",
        "CIV": "Civil Engineering",
        "CH": "Chemical Engineering",
        "BME": "Biomedical Engineering",
    }

    return course_dept_map.get(default_course, "Computer Science & Engineering")


class ParsedSchemeResponse(BaseModel):
    course_code: str | None = None
    theory_subjects: list[VTUSubject]
    practical_subjects: list[VTUSubject]


class FacultyItemModel(BaseModel):
    name: str
    department: str
    designation: str = "Assistant Professor"
    proficient_subjects: list[str] = []


class ParsedFacultyResponse(BaseModel):
    faculties: list[FacultyItemModel]


@router.get("/courses")
def get_vtu_courses():
    return PREFETCHED_VTU_COURSES


_EASYOCR_READER = None


def get_easyocr_reader():
    global _EASYOCR_READER
    if _EASYOCR_READER is None:
        import warnings
        warnings.filterwarnings("ignore")
        try:
            import torch
            import easyocr
            use_gpu = torch.cuda.is_available()
            _EASYOCR_READER = easyocr.Reader(['en'], gpu=use_gpu)
        except Exception:
            _EASYOCR_READER = None
    return _EASYOCR_READER


@router.post("/parse-scheme", response_model=ParsedSchemeResponse)
async def parse_vtu_scheme(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded")

    content = await file.read()
    extracted_text = ""

    ext = file.filename.lower().split(".")[-1]

    if ext == "pdf":
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            for page in doc:
                extracted_text += page.get_text() + "\n"
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to parse PDF document: {str(e)}")
    elif ext in ["png", "jpg", "jpeg", "webp", "bmp", "tiff"]:
        try:
            reader = get_easyocr_reader()
            if reader is not None:
                results = reader.readtext(content, detail=0)
                extracted_text = "\n".join(results)
        except Exception:
            pass

        if not extracted_text:
            try:
                from PIL import Image
                import pytesseract
                img = Image.open(io.BytesIO(content))
                extracted_text = pytesseract.image_to_string(img)
            except Exception:
                try:
                    doc = fitz.open(stream=content, filetype=ext)
                    for page in doc:
                        extracted_text += page.get_text() + "\n"
                except Exception:
                    extracted_text = content.decode("utf-8", errors="ignore")

    else:
        # Fallback text parsing for docx/txt
        try:
            extracted_text = content.decode("utf-8", errors="ignore")
        except Exception:
            extracted_text = str(content)

    theory_list: list[VTUSubject] = []
    practical_list: list[VTUSubject] = []

    lines = extracted_text.splitlines()
    vtu_code_regex = re.compile(r"\b(1[B]?[A-Z0-9]{2,8}[0-9]{2,3}[a-zA-Z]?|[0-9]{2}[A-Z]{2,6}[L]?[0-9]{2,3})\b")

    for line in lines:
        line_clean = line.strip()
        if not line_clean or "SCHEME" in line_clean.upper() or "VISVESVARAYA" in line_clean.upper():
            continue

        match = vtu_code_regex.search(line_clean)
        if match:
            code = match.group(1).upper()

            # Ignore common non-subject table header tokens
            if code in ["SEMESTER", "TEACHING", "QUESTION", "OUTCOME", "EXAMINATION"]:
                continue

            name_part = re.sub(r"^\d+\s+", "", line_clean)
            name_part = re.sub(r"\b" + re.escape(match.group(1)) + r"\b", "", name_part, flags=re.IGNORECASE)
            name_part = re.sub(r"\b(ASC|IPCC|PCC|PCCL|AEC|SDC|NCMC|TD\s*/?\s*PSB:[^\|]*)\b", "", name_part, flags=re.IGNORECASE)
            name_part = re.sub(r"[\d\.:\-\|\(\)]+", " ", name_part)
            name_part = " ".join(name_part.split()).strip()

            if not name_part or len(name_part) < 3:
                code_titles = {
                    "1BMATCS301": "Probability, Distributions and Statistics",
                    "1BCS302": "Object Oriented Programming with Java",
                    "1BCS303": "Digital Design and Computer Organization",
                    "1BCS304": "Operating Systems",
                    "1BCS305": "Data Structures and Applications",
                    "1BCSL306": "Data Structures Laboratory",
                    "1BXXL307X": "Ability Enhancement Course Lab",
                    "1BCP308": "Community Project / Societal Project",
                    "21CS32": "Data Structures and Applications",
                    "21CS33": "Analog and Digital Electronics",
                    "21CS34": "Computer Organization and Architecture",
                    "21CSL35": "Data Structures Laboratory",
                    "21CSL36": "Analog and Digital Electronics Laboratory",
                }
                name_part = code_titles.get(code, f"Subject {code}")

            is_lab = (
                "L" in code[3:]
                or "PCCL" in line_clean.upper()
                or "LAB" in line_clean.upper()
                or "LABORATORY" in line_clean.upper()
                or "PRACTICAL" in line_clean.upper()
                or "WORKSHOP" in line_clean.upper()
            )

            category = "practical" if is_lab else "theory"
            hours = 3 if category == "practical" else 4
            dept = infer_department(code, name_part)

            subj = VTUSubject(code=code, name=name_part, department=dept, category=category, weekly_hours=hours)

            if category == "practical":
                if not any(s.code == code for s in practical_list):
                    practical_list.append(subj)
            else:
                if not any(s.code == code for s in theory_list):
                    theory_list.append(subj)

    return ParsedSchemeResponse(
        course_code=None,
        theory_subjects=theory_list,
        practical_subjects=practical_list,
    )


def is_valid_human_name(name_str: str) -> bool:
    if not name_str or len(name_str) < 2 or len(name_str) > 80:
        return False
    # Disallow binary junk, zip entries, xml elements
    lower = name_str.lower()
    if any(bad in lower for bad in ["_rels", "[content_types]", ".xml", "<?xml", "pk\x03", "xmlns", "schemas."]):
        return False
    # Check for non-printable control characters
    if any(ord(c) < 32 and c not in "\t\n\r" for c in name_str):
        return False
    # Must contain mostly alphabetic characters
    letters = sum(1 for c in name_str if c.isalpha())
    if letters < 2:
        return False
    junk_chars = sum(1 for c in name_str if not (c.isalnum() or c.isspace() or c in ".,'-()&/"))
    if junk_chars > 3:
        return False
    return True


@router.post("/parse-faculty", response_model=ParsedFacultyResponse)
async def parse_vtu_faculty(file: UploadFile = File(...)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file uploaded")

    content = await file.read()
    ext = file.filename.lower().split(".")[-1]
    faculties: list[FacultyItemModel] = []

    # 1. Excel / CSV Parsing (xlsx, xls, csv, tsv)
    if ext in ["xlsx", "xls", "csv", "tsv"]:
        try:
            import pandas as pd
            if ext == "csv":
                try:
                    df = pd.read_csv(io.BytesIO(content))
                except Exception:
                    df = pd.read_csv(io.BytesIO(content), sep=None, engine="python")
            elif ext == "tsv":
                df = pd.read_csv(io.BytesIO(content), sep="\t")
            else:
                df = pd.read_excel(io.BytesIO(content))

            for _, row in df.iterrows():
                row_str = " ".join([str(val) for val in row.values if pd.notna(val)])
                if not row_str.strip():
                    continue

                val_name = ""
                val_dept = ""
                val_desg = "Assistant Professor"
                subjs: list[str] = []

                for col in df.columns:
                    col_l = str(col).lower().strip()
                    val = str(row[col]).strip() if pd.notna(row[col]) else ""
                    if not val or val.lower() in ["nan", "none", "null"]:
                        continue

                    if any(k in col_l for k in ["name", "faculty", "teacher", "professor", "instructor", "staff"]):
                        val_name = val
                    elif any(k in col_l for k in ["dept", "department", "branch", "stream"]):
                        val_dept = val
                    elif any(k in col_l for k in ["desig", "designation", "role", "post", "title", "position"]):
                        val_desg = val
                    elif any(k in col_l for k in ["subj", "course", "teach", "proficien", "handling", "allot"]):
                        subjs = [s.strip() for s in re.split(r"[,;/|]", val) if s.strip()]

                if not val_name:
                    for v in row.values:
                        if pd.notna(v):
                            v_str = str(v).strip()
                            if is_valid_human_name(v_str) and (re.search(r"\b(Dr|Prof|Mr|Mrs|Ms)\.?\b", v_str, re.I) or len(v_str.split()) <= 4):
                                val_name = v_str
                                break

                if not is_valid_human_name(val_name):
                    continue

                if not val_dept:
                    val_dept = infer_department(val_name + " " + row_str)

                # Clean designation
                if "prof" in val_name.lower() or "dr." in val_name.lower():
                    if val_desg == "Assistant Professor" and "assoc" in val_name.lower():
                        val_desg = "Associate Professor"
                    elif val_desg == "Assistant Professor" and "prof" in val_name.lower():
                        val_desg = "Professor"

                faculties.append(FacultyItemModel(
                    name=val_name,
                    department=val_dept,
                    designation=val_desg,
                    proficient_subjects=subjs,
                ))

            if faculties:
                return ParsedFacultyResponse(faculties=faculties)
        except Exception as err:
            print("Excel/CSV parse error:", err)

    # 2. Word Documents (.docx)
    elif ext in ["docx", "doc"]:
        try:
            import docx
            doc = docx.Document(io.BytesIO(content))
            
            # Check tables in Word document
            for table in doc.tables:
                headers = [cell.text.strip().lower() for cell in table.rows[0].cells] if len(table.rows) > 0 else []
                for row in table.rows[1:]:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = " ".join(cells)
                    if not row_text.strip():
                        continue
                    
                    val_name = ""
                    val_dept = ""
                    val_desg = "Assistant Professor"
                    subjs = []

                    for idx, cell_text in enumerate(cells):
                        col_hdr = headers[idx] if idx < len(headers) else ""
                        if any(k in col_hdr for k in ["name", "faculty", "staff", "teacher"]):
                            val_name = cell_text
                        elif any(k in col_hdr for k in ["dept", "department", "branch"]):
                            val_dept = cell_text
                        elif any(k in col_hdr for k in ["desig", "role", "position"]):
                            val_desg = cell_text
                        elif any(k in col_hdr for k in ["subj", "course"]):
                            subjs = [s.strip() for s in re.split(r"[,;/|]", cell_text) if s.strip()]

                    if not val_name:
                        for cell_text in cells:
                            if is_valid_human_name(cell_text) and len(cell_text.split()) <= 4:
                                val_name = cell_text
                                break

                    if is_valid_human_name(val_name):
                        if not val_dept:
                            val_dept = infer_department(val_name + " " + row_text)
                        faculties.append(FacultyItemModel(
                            name=val_name,
                            department=val_dept,
                            designation=val_desg,
                            proficient_subjects=subjs,
                        ))

            # Also check paragraphs
            if not faculties:
                for p in doc.paragraphs:
                    line_clean = p.text.strip()
                    if not line_clean or not is_valid_human_name(line_clean.split(",")[0]):
                        continue
                    parts = line_clean.split(",")
                    name = parts[0].strip()
                    dept = parts[1].strip() if len(parts) > 1 else infer_department(line_clean)
                    subjs_raw = parts[2].strip() if len(parts) > 2 else ""
                    subjs = [s.strip() for s in re.split(r"[;/|]", subjs_raw) if s.strip()]
                    desg = "Professor" if "Dr." in name or "Prof." in name else "Assistant Professor"
                    faculties.append(FacultyItemModel(
                        name=name,
                        department=dept,
                        designation=desg,
                        proficient_subjects=subjs,
                    ))

            if faculties:
                return ParsedFacultyResponse(faculties=faculties)
        except Exception as err:
            print("Docx parse error:", err)

    # 3. PDF Document Parsing
    elif ext == "pdf":
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            extracted_text = ""
            for page in doc:
                extracted_text += page.get_text() + "\n"

            for line in extracted_text.splitlines():
                line_clean = line.strip()
                if not line_clean or len(line_clean) < 3:
                    continue

                if re.search(r"\b(Dr|Prof|Mr|Mrs|Ms)\.?\b", line_clean, re.IGNORECASE) or (len(line_clean.split()) <= 4 and is_valid_human_name(line_clean)):
                    parts = line_clean.split(",")
                    name = parts[0].strip()
                    if not is_valid_human_name(name):
                        continue
                    dept = parts[1].strip() if len(parts) > 1 else infer_department(line_clean)
                    subjs_raw = parts[2].strip() if len(parts) > 2 else ""
                    subjs = [s.strip() for s in re.split(r"[;/|]", subjs_raw) if s.strip()]
                    desg = "Professor" if "Dr." in name or "Prof." in name else "Assistant Professor"

                    faculties.append(FacultyItemModel(
                        name=name,
                        department=dept,
                        designation=desg,
                        proficient_subjects=subjs,
                    ))
            if faculties:
                return ParsedFacultyResponse(faculties=faculties)
        except Exception as err:
            print("PDF parse error:", err)

    # 4. Text / OCR Parsing
    else:
        try:
            try:
                text = content.decode("utf-8")
            except UnicodeDecodeError:
                text = content.decode("latin1", errors="ignore")

            for line in text.splitlines():
                line_clean = line.strip()
                if not line_clean or not is_valid_human_name(line_clean.split(",")[0]):
                    continue

                parts = line_clean.split(",")
                name = parts[0].strip()
                dept = parts[1].strip() if len(parts) > 1 else infer_department(line_clean)
                subjs_raw = parts[2].strip() if len(parts) > 2 else ""
                subjs = [s.strip() for s in re.split(r"[;/|]", subjs_raw) if s.strip()]
                desg = "Professor" if "Dr." in name or "Prof." in name else "Assistant Professor"

                faculties.append(FacultyItemModel(
                    name=name,
                    department=dept,
                    designation=desg,
                    proficient_subjects=subjs,
                ))
            if faculties:
                return ParsedFacultyResponse(faculties=faculties)
        except Exception as err:
            print("Text parse error:", err)

    # 5. Return parsed faculties (or empty list)
    return ParsedFacultyResponse(faculties=faculties)

